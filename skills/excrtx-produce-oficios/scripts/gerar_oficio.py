#!/usr/bin/env python3
"""
Gerador de Ofícios — Gabinete Institucional
Gera DOCX, PDF ou HTML a partir de template + dados JSON.
"""
import argparse
import json
import os
import re
import sys
import zipfile
from datetime import date
from pathlib import Path

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False


def parse_template_frontmatter(template_path):
    """Extrai frontmatter YAML e corpo do template Markdown."""
    text = Path(template_path).read_text(encoding="utf-8")
    fm_match = re.match(r"^---\n(.*?)\n---\n(.*)$", text, re.DOTALL)
    if not fm_match:
        return {}, text

    if HAS_YAML:
        frontmatter = yaml.safe_load(fm_match.group(1))
    else:
        frontmatter = {}  # fallback: parse básico se necessário

    return frontmatter, fm_match.group(2)


def parse_frontmatter_basic(raw):
    """Parser YAML mínimo quando pyyaml não está disponível."""
    # Extrai campos_requeridos com regex
    campos = []
    match = re.search(r"campos_requeridos:\s*\n((?:\s+-.*\n|\s+\s+.*\n)*)", raw)
    if not match:
        return {"campos_requeridos": []}

    blocos = re.split(r"\n\s+- ", match.group(1))
    for bloco in blocos:
        campo = {}
        for line in bloco.strip().splitlines():
            line = line.strip().lstrip("- ")
            if ":" in line:
                key, val = line.split(":", 1)
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                if val == "true":
                    val = True
                elif val == "false":
                    val = False
                campo[key] = val
        if campo.get("nome"):
            campos.append(campo)

    return {"campos_requeridos": campos}


def get_fields(template_path):
    """Retorna lista de campos obrigatórios do template."""
    text = Path(template_path).read_text(encoding="utf-8")
    fm_match = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
    if not fm_match:
        return []

    raw = fm_match.group(1)
    if HAS_YAML:
        fm = yaml.safe_load(raw)
    else:
        fm = parse_frontmatter_basic(raw)

    return fm.get("campos_requeridos", [])


def fill_template(template_path, data):
    """Substitui placeholders {{campo}} no corpo do template."""
    _, corpo = parse_template_frontmatter(template_path)
    # Injeta dados automáticos
    data.setdefault("ano_atual", date.today().year)
    data.setdefault("localidade", "Localidade Padrão")
    data.setdefault("nome_signatario", "Lucas Vanini")
    data.setdefault("cargo_signatario", "Diretor Geral")

    for key, val in data.items():
        corpo = corpo.replace(f"{{{{{key}}}}}", str(val))

    return corpo


def corpo_markdown_to_paragrafos(texto):
    """Converte corpo markdown simples em lista de parágrafos estruturados."""
    paragrafos = []
    for linha in texto.strip().splitlines():
        linha = linha.rstrip()
        if not linha:
            paragrafos.append({"tipo": "blank", "texto": ""})
        elif linha.startswith("# "):
            paragrafos.append({"tipo": "titulo", "texto": linha[2:].strip()})
        elif linha.startswith("**") and linha.endswith("**"):
            paragrafos.append({"tipo": "texto", "texto": linha.strip("*"), "bold": True})
        elif "**" in linha:
            # Pode ter **Assunto:** texto
            paragrafos.append({"tipo": "texto", "texto": linha.replace("**", ""), "parse_bold": True, "original": linha})
        else:
            paragrafos.append({"tipo": "texto", "texto": linha})
    return paragrafos


def gerar_docx(template_path, data, output_path):
    """Gera DOCX usando python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Estilo padrão
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    corpo = fill_template(template_path, data)
    paragrafos = corpo_markdown_to_paragrafos(corpo)

    for p_info in paragrafos:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        if p_info["tipo"] == "blank":
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("")
            run.font.size = Pt(6)
            run.font.name = "Times New Roman"
            continue

        if p_info["tipo"] == "titulo":
            # Pular título principal (ofício) — já está no cabeçalho
            continue

        texto = p_info["texto"]

        if p_info.get("parse_bold") and "**" in p_info.get("original", ""):
            # Parse inline bold **...**
            original = p_info["original"]
            parts = re.split(r"(\*\*.*?\*\*)", original)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
        else:
            is_bold = p_info.get("bold", False)
            run = p.add_run(texto)
            run.bold = is_bold
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        # Alinhamento
        if p_info.get("bold"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Salvar DOCX
    docx_path = output_path if output_path.endswith(".docx") else output_path.replace(".zip", ".docx")
    doc.save(docx_path)

    # Se output termina em .zip, criar zip
    if output_path.endswith(".zip"):
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(docx_path, os.path.basename(docx_path))
        os.remove(docx_path)

    return output_path


def gerar_html(template_path, data, output_path):
    """Gera HTML estilizado."""
    corpo = fill_template(template_path, data)

    # Converter markdown simples para HTML
    html_corpo = ""
    for linha in corpo.strip().splitlines():
        if not linha.strip():
            html_corpo += "<br>\n"
        elif linha.startswith("# "):
            continue  # skip título
        elif "**" in linha:
            linha_html = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", linha)
            html_corpo += f"<p>{linha_html}</p>\n"
        else:
            html_corpo += f"<p>{linha}</p>\n"

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<title>Ofício {data.get('numero_oficio', '___')}/{data.get('ano_atual', date.today().year)}</title>
<style>
  @page {{
    margin: 3cm 2cm 2cm 3cm;
  }}
  body {{
    font-family: 'Times New Roman', 'Georgia', serif;
    font-size: 12pt;
    line-height: 1.5;
    color: #000;
    max-width: 700px;
    margin: 3cm auto;
    padding: 0 2cm;
  }}
  p {{
    text-align: justify;
    margin: 0 0 0.5em 0;
  }}
  strong {{
    font-weight: bold;
  }}
  .cabecalho {{
    text-align: center;
    margin-bottom: 2em;
  }}
  .cabecalho img {{
    max-height: 80px;
  }}
  .assinatura {{
    text-align: center;
    margin-top: 3em;
  }}
  .assinatura .nome {{
    font-weight: bold;
  }}
  @media print {{
    body {{
      margin: 0;
    }}
  }}
</style>
</head>
<body>
{html_corpo}
</body>
</html>"""

    Path(output_path).write_text(html, encoding="utf-8")
    return output_path


def gerar_pdf(template_path, data, output_path):
    """Gera PDF via HTML → WeasyPrint."""
    # Primeiro gerar HTML
    html_tmp = output_path.replace(".pdf", ".tmp.html")
    gerar_html(template_path, data, html_tmp)

    try:
        import weasyprint
        doc = weasyprint.HTML(filename=html_tmp)
        doc.write_pdf(output_path)
        os.remove(html_tmp)
        return output_path
    except ImportError:
        print("ERRO: weasyprint não instalado. Use: pip install weasyprint", file=sys.stderr)
        print("Alternativa: gerar em HTML e converter manualmente.", file=sys.stderr)
        # Limpar tmp
        if os.path.exists(html_tmp):
            os.remove(html_tmp)
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(description="Gerador de Ofícios Institucionais")
    parser.add_argument("--template", required=True, help="Caminho do template Markdown")
    parser.add_argument("--dados", required=True, help="JSON com dados do ofício")
    parser.add_argument("--formato", required=True, choices=["docx", "pdf", "html"],
                        help="Formato de saída")
    parser.add_argument("--output", required=True, help="Caminho do arquivo de saída")
    parser.add_argument("--listar-campos", action="store_true",
                        help="Listar campos obrigatórios e sair")

    args = parser.parse_args()

    if args.listar_campos:
        campos = get_fields(args.template)
        for c in campos:
            req = "★" if c.get("obrigatorio", True) else " "
            padrao = f" (padrão: {c.get('padrao', '')})" if c.get("padrao") else ""
            print(f"  {req} {c.get('label', c.get('nome'))} [{c.get('nome')}]{padrao}")
        return

    # Validar template
    if not os.path.exists(args.template):
        print(f"ERRO: template não encontrado: {args.template}", file=sys.stderr)
        sys.exit(1)

    # Parse dados
    try:
        data = json.loads(args.dados)
    except json.JSONDecodeError as e:
        print(f"ERRO: JSON inválido: {e}", file=sys.stderr)
        sys.exit(1)

    # Validar campos obrigatórios
    campos = get_fields(args.template)
    faltantes = []
    for campo in campos:
        nome = campo.get("nome")
        obrigatorio = campo.get("obrigatorio", True)
        tem_padrao = "padrao" in campo
        if obrigatorio and not tem_padrao and nome not in data:
            faltantes.append(campo.get("label", nome))

    if faltantes:
        print("ERRO: Campos obrigatórios faltantes:", file=sys.stderr)
        for f in faltantes:
            print(f"  - {f}", file=sys.stderr)
        sys.exit(1)

    # Gerar
    if args.formato == "docx":
        path = gerar_docx(args.template, data, args.output)
    elif args.formato == "pdf":
        path = gerar_pdf(args.template, data, args.output)
    elif args.formato == "html":
        path = gerar_html(args.template, data, args.output)

    size = os.path.getsize(path)
    print(f"Documento gerado: {path} ({size:,} bytes)")


if __name__ == "__main__":
    main()
